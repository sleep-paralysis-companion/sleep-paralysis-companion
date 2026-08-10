from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\satya\Documents\paralux\SP_App_Tech_Stack_and_Phase_1_Plan.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(92, 105, 122)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color: str = BORDER, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color: str = "7F9DB9", size: str = "12") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def set_run_font(run, size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_para(doc, text="", style=None, after=None) -> object:
    p = doc.add_paragraph(text, style=style)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_numbered(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_callout(doc, label: str, text: str, fill: str = CALLOUT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="D7DFEA", size="4")
    repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_simple_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent_dxa=120)
    set_table_borders(table)
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(value)
            set_run_font(r, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color, before, after in [
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        st = styles[list_style]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15


def add_masthead(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Sleep Paralysis Companion Technical Plan"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(hp.runs[0], size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(fp.add_run("Page "), size=9, color=MUTED)
    add_page_number(fp)

    add_para(doc, "Technical Recommendation & Phase 1 Build Plan").paragraph_format.space_after = Pt(2)
    title = doc.paragraphs[-1]
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.keep_with_next = True
    for run in title.runs:
        set_run_font(run, size=23, color=INK, bold=True)

    subtitle = add_para(
        doc,
        "Sleep Paralysis Companion - iOS-first MVP with Apple Watch and Android planned from day one",
    )
    subtitle.paragraph_format.space_after = Pt(14)
    for run in subtitle.runs:
        set_run_font(run, size=12.5, color=MUTED)

    rows = [
        ("Prepared for", "Client review"),
        ("Prepared by", "Satyam"),
        ("Date", "July 5, 2026"),
        ("Status", "Stack recommendation, delivery plan, and App Store readiness notes"),
    ]
    for label, value in rows:
        p = add_para(doc)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=10.5, bold=True, color=INK)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=INK)

    rule = add_para(doc)
    paragraph_border_bottom(rule)


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_masthead(doc)

    add_callout(
        doc,
        "My short recommendation",
        "I would build Phase 1 native on iOS using Swift and SwiftUI, with the codebase structured so Apple Watch, Android, Wear OS, HealthKit, Health Connect, subscriptions, and future AI/audio features can be added without rewriting the product. React Native is still an option if the highest priority is shipping phone UI on iOS and Android together, but for this PRD the hardest parts are platform-native: Apple Watch, notifications, lock screen actions, audio, HealthKit, StoreKit, privacy permissions, and App Store review.",
    )

    add_heading(doc, "How I Understand The Product", 1)
    add_para(
        doc,
        "The Phase 1 product is not just a basic journal. It is a sleep wellness app that helps a user prepare before sleep, get grounded after a sleep-paralysis episode, and build a simple personal pattern over time. The launch version should feel calm and trustworthy, but underneath it needs a serious architecture because later phases mention Apple Watch, Android watches, sleep-cycle mapping, AI reports, HealthKit/Health Connect, and possibly wearable biofeedback.",
    )
    add_para(
        doc,
        "For App Store safety, I would keep the product positioned as wellness, grounding, sleep preparation, and personal reflection. I would avoid diagnostic or treatment claims such as 'detects sleep paralysis', 'risk score', 'medical intervention', or anything that suggests the app is replacing a clinician. That language choice matters almost as much as the code.",
    )

    add_heading(doc, "1. React Native Vs Native iOS", 1)
    add_para(
        doc,
        "Both routes are possible. The choice depends on whether we optimize for shared phone UI now, or deeper Apple-platform quality first. My recommendation is native iOS first, but I am laying out the tradeoffs clearly so the decision is intentional.",
    )
    add_simple_table(
        doc,
        ["Option", "Pros", "Cons / Tradeoffs"],
        [
            [
                "React Native",
                "Good for shared iOS and Android phone screens. Faster for onboarding, forms, settings, journal screens, basic content flows, and teams already strong in JavaScript/TypeScript. Native modules can bridge into Swift/Kotlin when needed.",
                "Apple Watch still needs native SwiftUI. HealthKit, WidgetKit, App Intents, lock-screen actions, StoreKit, audio/background behavior, local encryption, and some permission flows usually become native bridge work. More dependency churn and more QA surface across OS releases.",
            ],
            [
                "Native iOS with Swift/SwiftUI",
                "Best fit for Apple Watch, HealthKit, WidgetKit, App Intents, AVFoundation audio, StoreKit 2, Keychain, CryptoKit, local notifications, accessibility, and App Store review. Cleaner path for a high-trust health/wellness app.",
                "Android will need a separate native app later. Less shared UI code. Requires strong Swift/iOS engineering discipline from the start.",
            ],
        ],
        [1800, 3720, 3840],
    )
    add_para(
        doc,
        "My call: I would choose native iOS for Phase 1. The reason is simple: the product's value is not just the screens. The value sits close to Apple's platform APIs and review rules. Starting native gives us better control, fewer compromises, and a more App Store-ready product.",
    )
    add_para(
        doc,
        "When React Native would make sense: if we decide Apple Watch is not part of the launch path, HealthKit is delayed, lock-screen/quick-action behavior is simplified, and the business goal becomes iOS + Android phone coverage as quickly as possible. That is a valid business choice, but it is not the best technical fit for the full PRD.",
    )

    add_heading(doc, "Recommended Tech Stack", 2)
    add_simple_table(
        doc,
        ["Layer", "Recommended Stack", "Why"],
        [
            ["iOS app", "Swift + SwiftUI", "Best Apple-platform fit and the cleanest route to App Store quality."],
            ["Architecture", "Modular clean architecture with MVVM/coordinators or reducer-style feature modules", "Keeps onboarding, audio, subscriptions, logs, and future watch/AI features separated."],
            ["Apple Watch", "watchOS + SwiftUI, planned from day one", "Even if the first release is iPhone-first, the app should be structured so a watch target can reuse domain logic and APIs."],
            ["Local storage", "SQLite with GRDB, or SwiftData if we intentionally target newer iOS versions only", "SQLite/GRDB is durable and portable; SwiftData is nice but more Apple-specific."],
            ["Audio", "AVFoundation", "Native playback, recording, sessions, interruption handling, and offline support."],
            ["Notifications", "UserNotifications locally, APNs for server-triggered flows later", "Bedtime and check-in reminders can start local; push can scale later."],
            ["Lock screen / quick action", "WidgetKit + App Intents", "The 'I just had an episode' entry point should be validated early because Apple controls what is possible here."],
            ["Subscriptions", "StoreKit 2", "Required path for in-app subscriptions and premium digital content on iOS."],
            ["Security", "Keychain, CryptoKit, Apple Data Protection, TLS", "Right baseline for a privacy-sensitive wellness product."],
            ["Android later", "Kotlin + Jetpack Compose; Wear OS native; Health Connect", "This avoids forcing watch and health integrations through a cross-platform compromise."],
        ],
        [1740, 3000, 4620],
    )

    add_heading(doc, "Backend Options", 2)
    add_para(
        doc,
        "There are two reasonable backend choices. I would personally choose the production AWS path if we want the strongest long-term compliance posture. If budget and speed matter more for MVP, Supabase/Firebase-class tooling can work, but I would keep the schema and APIs portable so we are not trapped later.",
    )
    add_simple_table(
        doc,
        ["Option", "What It Looks Like", "Pros", "Cons / Tradeoffs"],
        [
            [
                "Option A: AWS production stack",
                "TypeScript backend using NestJS or Fastify, PostgreSQL on RDS/Aurora, Redis, S3 + CloudFront for audio, queue workers, APNs integration, AWS KMS, Secrets Manager, CloudWatch, WAF.",
                "Best long-term control, stronger compliance posture, better auditability, easier to grow into AI/audio processing, B2B, insurer/employer, and health-data workflows.",
                "More setup and DevOps. Slightly slower than managed MVP platforms. Needs disciplined infrastructure and monitoring from the start.",
            ],
            [
                "Option B: Lean managed backend",
                "Supabase or Firebase-class backend. If choosing this route, I would lean Supabase because Postgres keeps the data model portable.",
                "Fastest to MVP, lower initial infrastructure cost, built-in auth/storage/functions, less operational work during build.",
                "Less control for complex compliance needs, future BAA/HIPAA posture needs careful validation, and advanced AI/audio pipelines may require migration or a hybrid backend later.",
            ],
        ],
        [1500, 2700, 2520, 2640],
    )
    add_para(
        doc,
        "My preference: native iOS app + AWS/Postgres backend if the goal is a serious, scalable, health-adjacent product. If the client wants the lowest-cost fastest MVP, I can start with a lean backend but would still design the API contracts, data model, and entitlement logic as if we may move to AWS later.",
    )

    add_heading(doc, "2. Scaling Plan And MVP Capacity", 1)
    add_para(
        doc,
        "For Phase 1, scale is very manageable if we do not upload continuous overnight audio recordings. The heavy media should be original audio content delivered through a CDN, while the app stores structured data: profile answers, episode logs, check-ins, consent events, subscription state, and optional voice clips if the user explicitly records or uploads them.",
    )
    add_bullets(
        doc,
        [
            "Stateless backend APIs so the app server can scale horizontally.",
            "PostgreSQL with correct indexes, migrations, backups, and read-scaling options.",
            "Object storage and CDN for audio files instead of serving audio from the API server.",
            "Queues for emails, push notifications, audio processing, analytics rollups, and future AI jobs.",
            "Redis for caching hot content, sessions/rate limits, and short-lived workflow state.",
            "Observability from the start: logs, metrics, error tracking, crash reporting, and alerting.",
            "Privacy-first analytics with no health data used for ads or marketing.",
        ]
    )
    add_callout(
        doc,
        "MVP scale estimate",
        "With AWS/Postgres/CDN and normal MVP traffic patterns, I would be comfortable designing Phase 1 for roughly 50k to 250k registered users, 10k to 50k monthly active users, and 2k to 10k daily active users before we need a major architecture change. A lean managed backend can also handle early MVP traction, but I would treat 10k to 50k MAU as the more comfortable planning band unless we tune and validate it. These numbers assume no continuous overnight audio uploads in Phase 1.",
        fill="EEF6FF",
    )
    add_para(
        doc,
        "If we add overnight audio recording/upload, the scaling story changes. Storage costs, background upload reliability, consent, App Store review notes, privacy labels, and data retention all become much more serious. I would keep continuous sleep recording out of the first App Store release unless it is absolutely required for launch.",
    )

    add_heading(doc, "3. Timeline To Build End To End", 1)
    add_para(
        doc,
        "With a focused team, I would plan this as a sub-8-week build. The goal should be to submit to App Store review by the end of week 7, while keeping week 8 for Apple review feedback, polish, and possible resubmission. Apple review timing is never fully under our control, so I would not promise a public App Store approval date as if it is purely an engineering task.",
    )
    add_simple_table(
        doc,
        ["Timing", "Focus", "Outcome"],
        [
            ["Week 1", "Scope lock, architecture, accounts, design system, data model, App Store risk review", "A buildable plan, repo setup, environments, and clarified launch scope."],
            ["Week 2", "Authentication, onboarding, profile questions, profile rules, base navigation", "User can sign up and complete the personalization flow."],
            ["Week 3", "Audio engine, offline audio catalog, bedtime alarm/audio flow, voice recording/upload path", "Core audio experience works on device."],
            ["Week 4", "Post-episode flow, lock-screen/quick-action spike, local notifications, morning check-in/log", "The main night and morning loops are usable."],
            ["Week 5", "Backend APIs, subscriptions, entitlement checks, privacy/account/data-rights screens", "App becomes commercially and operationally viable."],
            ["Week 6", "Apple Watch-ready architecture, optional basic watch companion, analytics, QA pass", "Watch path is de-risked and the MVP is close to complete."],
            ["Week 7", "TestFlight, bug fixes, App Store assets, privacy labels, review notes, production hardening", "Ready for App Store submission."],
            ["Week 8", "Apple review buffer, resubmission fixes, final polish", "Public release once approved by Apple."],
        ],
        [1320, 4320, 3720],
    )

    add_heading(doc, "4. Push Notifications And Payments", 1)
    add_heading(doc, "Push Notifications", 2)
    add_para(
        doc,
        "I would use local notifications first for bedtime reminders, morning check-ins, and simple in-app schedules. This is more reliable for MVP and does not require the server for every reminder. APNs push notifications should still be added in the backend foundation so we can later support server-triggered reminders, subscription lifecycle messaging, and partner/community features.",
    )
    add_para(
        doc,
        "We should also avoid making push permission mandatory. Apple does not like apps that block core functionality behind unnecessary permissions. The app should still work if the user says no to notifications, with clear prompts explaining what they gain if they opt in.",
    )
    add_heading(doc, "Payments", 2)
    add_callout(
        doc,
        "Important Apple Pay correction",
        "For in-app subscriptions, premium audio, AI reports, or locked digital features, we should use Apple In-App Purchase through StoreKit 2, not Apple Pay. Apple Pay is mainly for physical goods or services consumed outside the app. If the app sells digital functionality inside the app, StoreKit is the correct path.",
        fill="FFF8E8",
    )
    add_para(
        doc,
        "For Phase 1, I would implement StoreKit 2 with server-side entitlement verification so the app knows whether a user is trial, monthly, annual, lifetime, or expired. The planned pricing from the PRD can be set up as App Store products: 3-night trial, $8.99 monthly, about $59.99 yearly, $100 bi-yearly if Apple product configuration supports the exact structure cleanly, and $149 lifetime/founding member if we decide to keep that option.",
    )
    add_callout(
        doc,
        "Clarification needed on freemium",
        "After the 3-night trial ends, should users lose access completely until they subscribe, or should there be a free limited mode? This affects paywall copy, App Store review notes, onboarding, retention, and the data users can still access after trial expiry. My suggestion is a limited free mode: keep access to account, privacy controls, delete/export data, and a small amount of basic education/log viewing, while premium audio/programs/reports remain paid.",
    )

    add_heading(doc, "5. App Store Upload And Client Requirements", 1)
    add_para(
        doc,
        "I can help with the technical App Store path: certificates, bundle IDs, provisioning, TestFlight, app archive, App Store Connect setup, review notes, and the final upload. But a few items must come from the client because they are business/legal/account requirements.",
    )
    add_simple_table(
        doc,
        ["Client-side Requirement", "Why It Matters"],
        [
            ["Apple Developer Program account", "The app should be published under the client's Apple account, not mine. If it is an organization account, Apple may require business verification."],
            ["App Store Connect access", "I need the right role to create app records, add builds, configure subscriptions, and submit for review."],
            ["Paid apps agreement, tax, and banking", "Required before Apple can sell subscriptions or paid products."],
            ["Privacy policy URL and support URL", "Required for App Store submission and for a health/wellness product."],
            ["Final app name, brand assets, screenshots, and description approval", "Needed for metadata, screenshots, icon, and App Store product page."],
            ["Legal/compliance review", "Especially for consumer health data, consent, retention/deletion, and wellness-vs-medical claims."],
            ["Demo account or demo mode", "Apple Review needs full access to the app and subscriptions during review."],
        ],
        [3120, 6240],
    )
    add_para(
        doc,
        "For App Store positioning, I would submit this under Health & Fitness, not Medical. The app should be described as support, grounding, sleep wellness, and personal reflection. We should avoid any wording that implies diagnosis, treatment, medical monitoring, or automatic detection of sleep paralysis unless we later have the regulatory and clinical support for those claims.",
    )

    add_heading(doc, "6. Data Encryption And US Compliance Posture", 1)
    add_para(
        doc,
        "This is a sensitive wellness product, so I would design it with a higher bar than a normal lifestyle app. I would not call the first version HIPAA-compliant unless the legal/business setup actually requires HIPAA and the proper agreements are in place. But I would build it HIPAA-ready and consumer-health-data aware from day one.",
    )
    add_bullets(
        doc,
        [
            "On-device: store tokens and secrets in Keychain; protect local sensitive files with Apple Data Protection; encrypt voice files or sensitive local data where needed.",
            "In transit: TLS for all API calls and media downloads/uploads.",
            "Backend database: encrypted PostgreSQL storage using KMS-managed keys, encrypted backups, and restricted production access.",
            "Object storage: encrypted audio/voice assets using S3 server-side encryption, preferably SSE-KMS for sensitive files.",
            "Field-level protection: envelope encryption for especially sensitive values such as voice metadata, user health notes, and AI-report inputs if we store them.",
            "Data minimization: do not collect continuous microphone/audio data unless the feature is explicitly approved and consented.",
            "Consent and rights: clear consent screens, privacy settings, data export, delete account, delete voice recordings, and consent withdrawal.",
            "Access control: least-privilege IAM, admin audit logs, no raw production data in local dev, and strict separation between PII and wellness logs.",
            "No ad targeting with health data: health/wellness data should not be used for advertising, marketing, or data mining.",
        ]
    )
    add_para(
        doc,
        "For US privacy, I would prepare for consumer health data obligations, not just generic privacy. That means a clear consumer health data privacy policy, retention/deletion language, consent logs, and a product posture that respects state-level health privacy expectations. Final legal wording should be reviewed by counsel.",
    )

    add_heading(doc, "7. AI-Assisted Development / Vibe Coding", 1)
    add_para(
        doc,
        "I do use AI-assisted development, but I do not treat it as a replacement for engineering judgment. I use multiple harnesses depending on the task: Codex, Claude Code, OpenCode, and Antigravity, along with agent skills and specific models for specific workstreams like architecture, Swift implementation, UI polish, security review, testing, documentation, and refactoring.",
    )
    add_para(
        doc,
        "My workflow is: use AI to move faster on scaffolding, repetitive implementation, test generation, and review passes; then manually inspect the critical parts myself. Anything involving payments, privacy, encryption, App Store policy, authentication, health data, or production infrastructure gets human review and test coverage before it is treated as done.",
    )

    add_heading(doc, "A Few Things I Would Add To The Scope", 1)
    add_bullets(
        doc,
        [
            "App Store wording review before development goes too far, so we do not build a feature that later creates review or regulatory friction.",
            "A technical spike for the lock-screen/quick-action flow in week 1 or 2, because Apple controls what is possible there.",
            "A clear decision on whether Phase 1 includes any overnight microphone/sleep recording. I recommend not including continuous recording in the first public release.",
            "Privacy and consent design as product features, not paperwork added at the end.",
            "A real TestFlight period with physical-device testing before App Store submission.",
            "Subscription entitlement logic on the backend, so future Android purchases can map into one account system.",
        ]
    )

    doc.add_page_break()
    add_heading(doc, "Phase 1 Milestone Plan", 1)
    add_para(
        doc,
        "Below is how I would build Phase 1 end to end. The plan keeps the client's expectations intact, but adds the engineering details needed to overdeliver: cleaner architecture, App Store readiness, privacy screens, proper subscription handling, and a watch-ready foundation.",
    )

    milestones = [
        [
            "M0",
            "Scope, Risk, And Launch Alignment",
            "Confirm Phase 1 feature list; lock wellness language; choose native iOS/backend option; confirm pricing; clarify freemium; decide on sleep recording; collect Apple account requirements.",
            "Signed-off launch scope, final tech-stack direction, App Store risk notes, initial data map.",
        ],
        [
            "M1",
            "Project Foundation",
            "Create iOS repo structure; set modular architecture; configure environments; create design system; set local persistence; backend scaffold; CI; crash/error tooling.",
            "A clean, scalable codebase that the team can build on without rewrites.",
        ],
        [
            "M2",
            "Onboarding And Personalization",
            "Splash, sign up with Apple/email, explainer, three onboarding questions, profile rules, 'what we set up for you' screen, consent-aware first-run flow.",
            "User can onboard and receive a personalized setup path.",
        ],
        [
            "M3",
            "Audio And Voice Experience",
            "Long pre-sleep audio, short grounding audio, offline audio handling, audio session interruptions, user voice recording, partner voice upload/share script, delete voice recording.",
            "The emotional core of the product works reliably on-device.",
        ],
        [
            "M4",
            "Night And Morning Loops",
            "Set alarm, bedtime flow, local notifications, lock-screen/quick-action spike, post-episode grounding screen, call partner/hear voice/play calming audio, morning check-in/log.",
            "A complete first-night-to-morning user journey.",
        ],
        [
            "M5",
            "Subscriptions And Entitlements",
            "StoreKit 2 products, 3-night trial, monthly/yearly/bi-yearly/lifetime products as finalized, paywall, restore purchases, backend entitlement verification.",
            "Paid product is ready for App Store review and real customers.",
        ],
        [
            "M6",
            "Privacy, Compliance, And Account Controls",
            "Privacy settings, data rights, manage consent, delete account, export data, delete voice data, AI disclosure placeholder if AI reports are included, retention settings.",
            "A privacy-sensitive app that feels trustworthy and review-ready.",
        ],
        [
            "M7",
            "Watch-Ready Foundation",
            "Create shared domain layer for future watch app; validate watch communication assumptions; optionally ship a basic watch companion/quick action if schedule allows.",
            "Apple Watch path is de-risked instead of becoming a Phase 2 rewrite.",
        ],
        [
            "M8",
            "QA, TestFlight, And App Store Submission",
            "Device testing, accessibility pass, subscription sandbox testing, privacy labels, review notes, screenshots, metadata, demo account, final archive and upload.",
            "Submitted build that can realistically pass App Store review.",
        ],
    ]
    add_simple_table(doc, ["Milestone", "Focus", "Build Work", "Definition Of Done"], milestones, [780, 1920, 3900, 2760])

    add_heading(doc, "Phase 1 Deliverables", 2)
    add_bullets(
        doc,
        [
            "Native iOS app built in Swift/SwiftUI.",
            "Scalable app architecture with clean module boundaries.",
            "Sign up, onboarding, profile logic, and personalized setup flow.",
            "Audio playback and voice-recording experience.",
            "Bedtime alarm/audio flow, post-episode grounding flow, and morning check-in/log.",
            "Local notifications and APNs-ready backend foundation.",
            "StoreKit 2 subscriptions/trial/lifetime entitlement support.",
            "Privacy, consent, data export, delete account, and voice deletion controls.",
            "Backend APIs, database, object storage, admin-safe operational setup.",
            "TestFlight build and App Store submission support.",
        ]
    )

    add_heading(doc, "What I Would Not Put In The First Public Release", 2)
    add_para(
        doc,
        "To keep the first release strong and App Store-safe, I would avoid a few tempting but risky items until the product has user feedback and the compliance path is clearer.",
    )
    add_bullets(
        doc,
        [
            "Continuous overnight microphone recording by default.",
            "Claims that the app detects, diagnoses, predicts, or treats sleep paralysis.",
            "AI reports that sound medical or clinical without review.",
            "Community features without moderation tooling.",
            "Complex wearable biofeedback alerts before the basic Apple Watch path is proven.",
        ]
    )

    add_heading(doc, "Decision Points For Client", 1)
    add_simple_table(
        doc,
        ["Decision", "My Suggested Direction", "Why"],
        [
            ["Mobile stack", "Native iOS first", "Better for Apple Watch, HealthKit, StoreKit, audio, notifications, and App Store review."],
            ["Backend", "AWS/Postgres if budget allows; Supabase-class lean backend if speed/cost is the priority", "Both can work. AWS is stronger for long-term compliance and scale."],
            ["Payment method", "StoreKit 2, not Apple Pay", "Apple requires in-app purchase for digital subscriptions and premium in-app features."],
            ["Freemium", "Limited free mode after 3-night trial", "Better user trust and smoother data-access/privacy behavior after trial expiry."],
            ["Sleep recording", "Do not include continuous overnight recording in first public MVP", "Reduces App Store, privacy, storage, and compliance risk."],
            ["Apple Watch", "Build watch-ready now; basic watch companion only if schedule allows", "Keeps us future-proof without risking the iOS MVP timeline."],
        ],
        [1920, 3600, 3840],
    )

    add_heading(doc, "Reference Links Reviewed", 1)
    add_para(
        doc,
        "These are the primary references I would use while implementing and preparing the App Store submission. They should be checked again near submission because platform rules can change.",
    )
    refs = [
        ("Apple App Review Guidelines", "https://developer.apple.com/app-store/review/guidelines/"),
        ("Apple App Privacy Details", "https://developer.apple.com/app-store/app-privacy-details/"),
        ("Apple Health and Fitness", "https://developer.apple.com/health-fitness/"),
        ("Apple watchOS developer overview", "https://developer.apple.com/watchos/"),
        ("React Native native platform docs", "https://reactnative.dev/docs/native-platform"),
        ("Android Health Connect", "https://developer.android.com/health-and-fitness/health-connect"),
        ("Google Play Billing", "https://developer.android.com/google/play/billing"),
        ("AWS HIPAA Compliance", "https://aws.amazon.com/compliance/hipaa-compliance/"),
        ("AWS RDS encryption", "https://docs.aws.amazon.com/AmazonRDS/latest/UserGuide/Overview.Encryption.html"),
        ("AWS S3 server-side encryption", "https://docs.aws.amazon.com/AmazonS3/latest/userguide/UsingServerSideEncryption.html"),
        ("AWS KMS overview", "https://docs.aws.amazon.com/kms/latest/developerguide/overview.html"),
        ("Apple Keychain data protection", "https://support.apple.com/guide/security/keychain-data-protection-secb0694df1a/web"),
    ]
    for name, url in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        add_hyperlink(p, name, url)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
